"""
DOCX Document Generation

Provides professional DOCX generation capabilities for reports and documents.
Supports comprehensive formatting, styles, and Microsoft Word compatibility.

Author: Luke Steuber
"""

import re
import logging
from datetime import datetime
from typing import Dict, List, Optional, Any
from pathlib import Path

logger = logging.getLogger(__name__)

# DOCX Generation imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX generation disabled")


class DOCXGenerator:
    """Generate professional DOCX documents from structured content"""

    def __init__(
        self,
        output_dir: str = "reports/docx",
        title_color: tuple = (31, 78, 121),  # RGB
        default_font: str = "Arial"
    ):
        """
        Initialize DOCX generator

        Args:
            output_dir: Directory for DOCX output
            title_color: RGB tuple for title color (default: blue)
            default_font: Default font family
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with: pip install python-docx"
            )

        self.title_color = RGBColor(*title_color)
        self.default_font = default_font

    def generate_report_docx(
        self,
        content_sections: List[Dict[str, Any]],
        title: str,
        document_id: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Generate DOCX report from structured content sections

        Args:
            content_sections: List of content sections, each with:
                - title: Section title
                - content: Section content (markdown-style string)
                - type: Optional section type
            title: Document title
            document_id: Unique document identifier
            metadata: Optional metadata dict

        Returns:
            Dict with success status, filepath, filename, and size
        """
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"report_{document_id}_{timestamp}.docx"
            filepath = self.output_dir / filename

            # Create document
            doc = Document()
            self._setup_document_styles(doc)

            # Title page
            title_para = doc.add_paragraph(title, style='DocumentTitle')
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacing

            # Metadata section
            if metadata:
                doc.add_heading('Document Information', level=1)

                meta_para = doc.add_paragraph()
                meta_para.add_run(f"Document ID: {document_id}").bold = True
                meta_para.add_run(
                    f"\nGenerated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
                )

                # Add custom metadata
                for key, value in metadata.items():
                    formatted_key = key.replace('_', ' ').title()
                    meta_para.add_run(f"\n{formatted_key}: {value}")

                doc.add_page_break()

            # Content sections
            for i, section in enumerate(content_sections):
                section_title = section.get('title', f'Section {i+1}')
                section_content = section.get('content', '')

                # Section header
                doc.add_heading(section_title, level=1)

                # Section content
                self._add_docx_content(doc, section_content)

                # Page break between major sections
                if i < len(content_sections) - 1:
                    doc.add_page_break()

            # Save document
            doc.save(str(filepath))

            return {
                "success": True,
                "filepath": str(filepath),
                "filename": filename,
                "size_bytes": filepath.stat().st_size
            }

        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return {"success": False, "error": str(e)}

    def _setup_document_styles(self, doc: "Document"):
        """Set up custom document styles"""
        styles = doc.styles

        # Document title style
        try:
            title_style = styles.add_style('DocumentTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = self.default_font
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = self.title_color
        except Exception:
            # Style might already exist
            pass

        # Heading styles (customize if needed)
        try:
            heading1 = styles['Heading 1']
            heading1.font.color.rgb = self.title_color
            heading1.font.size = Pt(18)
        except Exception:
            pass

    def _add_docx_content(self, doc: "Document", content: str):
        """
        Add formatted content to DOCX document

        Supports markdown-style formatting:
        - # Heading 1
        - ## Heading 2
        - ### Heading 3
        - - List item
        - ``` code block
        """
        lines = content.split('\n')
        in_code_block = False
        code_lines = []

        for line in lines:
            line_stripped = line.strip()

            # Handle code blocks
            if line_stripped.startswith('```'):
                if in_code_block:
                    # End code block - add as pre-formatted text
                    if code_lines:
                        code_para = doc.add_paragraph()
                        code_run = code_para.add_run('\n'.join(code_lines))
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(10)
                        # Light gray background (approximation)
                        code_para.paragraph_format.left_indent = Inches(0.5)
                    code_lines = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                    # Skip language specifier if present
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # Empty lines
            if not line_stripped:
                doc.add_paragraph()
                continue

            # Markdown headings
            if line_stripped.startswith('# '):
                doc.add_heading(line_stripped[2:], level=1)
            elif line_stripped.startswith('## '):
                doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith('### '):
                doc.add_heading(line_stripped[4:], level=3)
            # Lists
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                doc.add_paragraph(line_stripped[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line_stripped):
                # Numbered list
                item_text = re.sub(r'^\d+\.\s', '', line_stripped)
                doc.add_paragraph(item_text, style='List Number')
            # Bold text **text**
            elif '**' in line_stripped:
                para = doc.add_paragraph()
                self._add_formatted_run(para, line_stripped)
            # Regular paragraphs
            else:
                doc.add_paragraph(line_stripped)

    def _add_formatted_run(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic)"""
        # Simple bold parsing
        parts = re.split(r'(\*\*.*?\*\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                # Normal text
                paragraph.add_run(part)


def generate_docx_report(
    content_sections: List[Dict[str, Any]],
    title: str,
    document_id: str,
    output_dir: str = "reports/docx",
    metadata: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Convenience function to generate a DOCX report

    Args:
        content_sections: List of content sections
        title: Document title
        document_id: Unique document identifier
        output_dir: Output directory
        metadata: Optional metadata

    Returns:
        Dict with success status and file information
    """
    try:
        generator = DOCXGenerator(output_dir=output_dir)
        return generator.generate_report_docx(
            content_sections=content_sections,
            title=title,
            document_id=document_id,
            metadata=metadata
        )
    except Exception as e:
        logger.error(f"DOCX report generation failed: {e}")
        return {"success": False, "error": str(e)}
